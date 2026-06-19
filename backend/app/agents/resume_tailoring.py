"""
CareerPilot AI — Agent 4: Resume Tailoring Agent
=================================================
Optimizes resume content to match specific job descriptions
without fabricating any qualifications.

Strict rules:
  - Never add skills/experience/education not in original resume
  - Safety check via extract_nouns() to compare original vs tailored
  - PDF generation with WeasyPrint
  - DOCX generation with python-docx
  - Output storage to ResumeVersion linked to JobPosting
"""

import json
import logging
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

from groq import Groq

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Safety helpers
# ---------------------------------------------------------------------------

TECH_TERMS: Set[str] = {
    "AWS", "Azure", "GCP", "Docker", "K8s", "Kubernetes",
    "Terraform", "Ansible", "Jenkins", "Linux", "Python",
    "Bash", "Git", "CI/CD", "DevOps", "SRE", "IaC", "Pulumi",
    "Prometheus", "Grafana", "ELK", "EFK", "Fluentd", "Istio",
    "Helm", "ArgoCD", "GitOps", "Vault", "Consul", "Nomad",
    "Packer", "Vagrant", "CircleCI", "GitHub", "GitLab",
    "Bitbucket", "Jira", "Confluence", "Slack", "Splunk",
    "Datadog", "NewRelic", "CloudWatch", "Lambda", "ECS",
    "EKS", "S3", "EC2", "RDS", "VPC", "IAM", "CloudFormation",
    "CDK", "OpenStack", "VMware", "VirtualBox", "MySQL",
    "PostgreSQL", "MongoDB", "Redis", "Kafka", "RabbitMQ",
    "Nginx", "Apache", "HAProxy", "Traefik", "Caddy",
    "Systemd", "Supervisor", "SELinux", "AppArmor",
}


def extract_nouns(text: str) -> Set[str]:
    """Extract proper nouns and known tech terms from *text*.

    Used to compare original vs tailored resume content to detect
    fabricated qualifications.
    """
    # Capitalised words that look like proper nouns / brand names
    capitalised = set(re.findall(r"\b[A-Z][a-zA-Z0-9+#.-]{1,}\b", text))
    # Tech terms (case-insensitive match but return original casing)
    tech_found = set()
    for term in TECH_TERMS:
        if re.search(rf"\b{re.escape(term)}\b", text, re.IGNORECASE):
            tech_found.add(term)
    return capitalised | tech_found


def safety_check(tailored: Dict[str, Any], original: Dict[str, Any]) -> List[str]:
    """Compare tailored resume against original to detect fabrications.

    Returns a list of warning messages.  An empty list means the check
    passed.
    """
    original_text = json.dumps(original)
    tailored_text = json.dumps(tailored)

    original_nouns = extract_nouns(original_text)
    tailored_nouns = extract_nouns(tailored_text)

    new_items = tailored_nouns - original_nouns
    warnings: List[str] = []
    if new_items:
        warnings.append(
            f"Potential fabrication detected — new nouns not in original: "
            f"{sorted(new_items)}"
        )
    return warnings


# ---------------------------------------------------------------------------
# Resume Tailor
# ---------------------------------------------------------------------------

class ResumeTailor:
    """Agent 4: Tailor a resume to a specific job description.

    Usage::

        tailor = ResumeTailor()
        result = tailor.tailor_resume(original_resume_dict, job_description)
        tailor.generate_pdf(result, "output.pdf")
        tailor.generate_docx(result, "output.docx")
    """

    MODEL = "llama-3.3-70b-versatile"

    def __init__(self, api_key: Optional[str] = None) -> None:
        self.client = Groq(api_key=api_key or os.getenv("GROQ_API_KEY"))
        if not self.client.api_key:
            raise ValueError(
                "GROQ_API_KEY must be set via environment variable or "
                "passed as api_key argument."
            )

    # -- Tailoring ----------------------------------------------------------

    def tailor_resume(
        self,
        original_resume: Dict[str, Any],
        job_description: str,
    ) -> Dict[str, Any]:
        """Run the LLM tailoring pass and return the optimised resume JSON.

        Raises ``ValueError`` if the safety check fails (i.e. the model
        added fabricated content).
        """
        system_prompt = (
            "You are an expert resume optimizer. You will receive:\n"
            "1. The user's original resume (JSON format)\n"
            "2. A job description\n\n"
            "Your task: Rewrite the resume to better match the job "
            "description.\n\n"
            "HARD RULES — violating these is not allowed:\n"
            "- Do NOT add any skill, tool, technology, certification, "
            "or project not present in the original resume.\n"
            "- Do NOT modify employment dates, company names, or job "
            "titles.\n"
            "- Do NOT fabricate any achievement or metric.\n"
            "- You MAY reorder content, improve phrasing, and surface "
            "relevant keywords.\n\n"
            "Return ONLY the optimized resume JSON.  No explanations."
        )

        user_prompt = (
            f"Original Resume:\n{json.dumps(original_resume, indent=2)}\n\n"
            f"Job Description:\n{job_description}"
        )

        logger.info("Sending tailoring request to Groq (%s)", self.MODEL)
        response = self.client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            model=self.MODEL,
            temperature=0.2,
            response_format={"type": "json_object"},
        )

        tailored: Dict[str, Any] = json.loads(
            response.choices[0].message.content
        )

        # Safety check — detect fabrications
        warnings = safety_check(tailored, original_resume)
        if warnings:
            warning_text = "; ".join(warnings)
            logger.warning("Safety check failed: %s", warning_text)
            raise ValueError(
                f"Tailoring safety check failed. {warning_text}"
            )

        logger.info("Tailoring completed successfully — safety check passed.")
        return tailored

    # -- PDF generation -----------------------------------------------------

    @staticmethod
    def generate_pdf(
        resume_json: Dict[str, Any],
        output_path: str,
    ) -> str:
        """Render the resume as a PDF via WeasyPrint."""
        full_name = resume_json.get("full_name", resume_json.get("name", "Resume"))
        summary = resume_json.get("summary", "")

        experience_html = ""
        for exp in resume_json.get("experience", []):
            role = exp.get("role", exp.get("title", ""))
            company = exp.get("company", "")
            duration = exp.get("duration", "")
            responsibilities = exp.get("responsibilities", [])
            items = "".join(
                f"<li>{r}</li>" for r in responsibilities
            )
            experience_html += f"""
            <div class="experience-block">
                <h3>{role} at {company}</h3>
                <p class="duration">{duration}</p>
                <ul>{items}</ul>
            </div>
            """

        education_html = ""
        for edu in resume_json.get("education", []):
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            education_html += (
                f"<p><strong>{degree}</strong> — {institution} ({year})</p>"
            )

        skills_html = ""
        skills = resume_json.get("skills", [])
        if isinstance(skills, list):
            skills_html = "".join(
                f"<span class='skill'>{s}</span> " for s in skills
            )

        html_content = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    body {{ font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 11pt;
           color: #222; margin: 40px; line-height: 1.5; }}
    h1 {{ font-size: 22pt; margin-bottom: 4px; }}
    .contact {{ color: #555; margin-bottom: 16px; }}
    h2 {{ font-size: 14pt; border-bottom: 2px solid #1a73e8;
          padding-bottom: 4px; margin-top: 20px; }}
    .experience-block {{ margin-bottom: 14px; }}
    .experience-block h3 {{ font-size: 12pt; margin-bottom: 2px; }}
    .duration {{ color: #777; font-style: italic; margin: 2px 0 6px; }}
    ul {{ margin: 4px 0 0 18px; padding: 0; }}
    li {{ margin-bottom: 3px; }}
    .skill {{ display: inline-block; background: #e8f0fe; color: #1a73e8;
              padding: 2px 8px; border-radius: 3px; margin: 2px 3px;
              font-size: 10pt; }}
    .summary {{ margin: 10px 0; }}
</style>
</head>
<body>
    <h1>{full_name}</h1>
    <div class="contact">{resume_json.get('contact_email', resume_json.get('email', ''))}</div>
    <div class="summary"><p>{summary}</p></div>

    <h2>Experience</h2>
    {experience_html}

    <h2>Education</h2>
    {education_html}

    <h2>Skills</h2>
    <p>{skills_html}</p>
</body>
</html>"""

        output_path = str(Path(output_path).resolve())
        from weasyprint import HTML  # type: ignore[import-untyped]
        HTML(string=html_content).write_pdf(output_path)
        logger.info("PDF generated: %s", output_path)
        return output_path

    # -- DOCX generation ----------------------------------------------------

    @staticmethod
    def generate_docx(
        resume_json: Dict[str, Any],
        output_path: str,
    ) -> str:
        """Render the resume as a .docx file via python-docx."""
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt, Inches  # type: ignore[import-untyped]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        full_name = resume_json.get("full_name", resume_json.get("name", "Resume"))

        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = name_para.add_run(full_name)
        run.bold = True
        run.font.size = Pt(22)

        # Contact
        email = resume_json.get("contact_email", resume_json.get("email", ""))
        if email:
            doc.add_paragraph(email)

        # Summary
        summary = resume_json.get("summary", "")
        if summary:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(summary)

        # Experience
        for exp in resume_json.get("experience", []):
            role = exp.get("role", exp.get("title", ""))
            company = exp.get("company", "")
            duration = exp.get("duration", "")
            heading_text = f"{role} at {company}"
            if duration:
                heading_text += f" ({duration})"
            doc.add_heading(heading_text, level=3)

            for resp in exp.get("responsibilities", []):
                doc.add_paragraph(resp, style="List Bullet")

        # Education
        for edu in resume_json.get("education", []):
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            doc.add_paragraph(f"{degree} — {institution} ({year})")

        # Skills
        skills = resume_json.get("skills", [])
        if isinstance(skills, list) and skills:
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(", ".join(skills))

        output_path = str(Path(output_path).resolve())
        doc.save(output_path)
        logger.info("DOCX generated: %s", output_path)
        return output_path

    # -- Persistence helpers ------------------------------------------------

    @staticmethod
    def build_resume_version_data(
        tailored_json: Dict[str, Any],
        original_resume_id: str,
        job_posting_id: str,
    ) -> Dict[str, Any]:
        """Build a dict suitable for inserting into a ``resume_versions``
        database row.

        Callers are responsible for the actual DB insert (SQLAlchemy / raw
        asyncpg).  Example::

            data = tailor.build_resume_version_data(
                tailored_json, original_id, job_id
            )
            await db.execute(
                resume_versions.insert().values(**data)
            )
        """
        return {
            "resume_id": original_resume_id,
            "job_posting_id": job_posting_id,
            "version_data": json.dumps(tailored_json),
            "version_type": "tailored",
            "created_at": None,  # let DB server_default handle it
        }


# ---------------------------------------------------------------------------
# Convenience entrypoint
# ---------------------------------------------------------------------------

async def run_resume_tailoring_task(
    original_resume: Dict[str, Any],
    job_description: str,
    original_resume_id: str,
    job_posting_id: str,
    output_dir: Optional[str] = None,
) -> Dict[str, Any]:
    """High-level convenience: tailor, generate PDF+docx, return metadata.

    Designed to be called from a Celery task or FastAPI endpoint.

    Returns a dict with keys:
        tailored_json, pdf_path, docx_path, version_data, warnings
    """
    tailor = ResumeTailor()
    tailored = tailor.tailor_resume(original_resume, job_description)

    if output_dir is None:
        output_dir = os.getenv(
            "CAREERPILOT_STORAGE",
            "/app/storage/resume_versions",
        )
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    pdf_path = str(out / f"tailored_{job_posting_id}.pdf")
    docx_path = str(out / f"tailored_{job_posting_id}.docx")

    tailor.generate_pdf(tailored, pdf_path)
    tailor.generate_docx(tailored, docx_path)

    version_data = tailor.build_resume_version_data(
        tailored, original_resume_id, job_posting_id
    )

    return {
        "tailored_json": tailored,
        "pdf_path": pdf_path,
        "docx_path": docx_path,
        "version_data": version_data,
        "warnings": [],
    }
