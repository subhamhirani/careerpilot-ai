from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent
CONTACT = "Ahmedabad, India | subhamhirani001@gmail.com | +91 98750 27571 | linkedin.com/in/subham-hirani | github.com/subhamhirani"
BASE_SKILLS = "Networking: TCP/IP, DNS, DHCP, VLANs, Routing & Switching, OSPF, NAT/PAT, VPN, RRAS, ACLs, Wireshark, STP | Windows Server: AD DS, GPO, DNS/DHCP, Hyper-V, IIS | Linux: Ubuntu, CentOS/RHEL, SSH, Bash | Infrastructure: Docker, GitLab CE, Git, Restic, backup/DR, PowerShell, Uptime Kuma | AWS: EC2, VPC, IAM, S3"
EXPERIENCE = [
    ("Network & Infrastructure Engineer | AnaxisTech LLP | Ahmedabad | Apr 2026–Present", [
        "Deployed self-hosted GitLab CE on Docker with persistent volumes, reverse proxy, authentication, and tested backup/restore procedures.",
        "Configured a multi-ISP redundant RRAS network using static public IPs and port forwarding, removing a single point of failure.",
        "Installed and administered Windows Server 2025 with AD DS, DNS, DHCP, and GPO security baselines.",
        "Built centralized automated backup and disaster-recovery workflows with Windows Task Scheduler, execution logging, and integrity verification.",
        "Troubleshot LAN/WAN, IIS FTP, SSH/SFTP, and Windows networking incidents; maintained topology diagrams, firewall matrices, and runbooks.",
    ]),
    ("Networking & Cybersecurity Trainee | Tops Technologies | Ahmedabad | Dec 2024–Apr 2026", [
        "Completed full-time practical networking and security training covering CCNA routing and switching, subnetting, VLANs, OSPF, ACLs, NAT/PAT, and STP.",
        "Provisioned AD DS, DNS, DHCP, Hyper-V, and Group Policy lab environments; administered Ubuntu, CentOS/RHEL, and AWS fundamentals.",
        "Built a three-VM Wazuh SIEM lab, simulated brute-force and port-scan events, tuned detections, and documented incident triage.",
    ]),
]
JOBS = [
    {
        "slug": "01_system_administrator_gandhinagar",
        "title": "System Administrator – Windows/Linux Server",
        "company": "Big Ideas Social Media Recruitment",
        "location": "Gandhinagar, Gujarat",
        "url": "https://www.shine.com/jobs/system-administrator-windows-linux-server/big-ideas-social-media-recruitment/19188395",
        "summary": "Network & Infrastructure Engineer with production Windows Server, Active Directory, DNS/DHCP, Hyper-V, Linux administration, Docker, backup/DR, and PowerShell/Bash experience. Proven in server operations, infrastructure troubleshooting, automation, security hardening, and documentation.",
        "focus": ["Windows Server 2025, AD DS, DNS, DHCP, GPO, Hyper-V", "Linux administration, Bash, PowerShell, service management", "Backup/DR automation, Docker, GitLab CE, documentation"],
        "letter": "Your Windows/Linux infrastructure role closely matches my production work at AnaxisTech LLP. I administer Windows Server 2025 services including AD DS, DNS, DHCP, GPO, and Hyper-V; support Linux and Docker-hosted services; and build automated backup workflows with verification and logging. My full-time networking and security training further strengthened my troubleshooting, routing, and infrastructure-security fundamentals. I would bring hands-on server operations, documentation discipline, and a rapid incident-resolution mindset to your Gandhinagar team.",
    },
    {
        "slug": "02_support_system_engineer_gandhinagar",
        "title": "Support System Engineer – Windows/Linux Server",
        "company": "Big Ideas Social Media Recruitment",
        "location": "Gandhinagar, Gujarat",
        "url": "https://www.shine.com/jobs/support-system-engineer-windows-linux-server/big-ideas-social-media-recruitment/19185040",
        "summary": "Infrastructure support engineer with hands-on Windows Server, Linux, Active Directory, DNS/DHCP, Docker, monitoring, backup/DR, and network troubleshooting experience in production and full-time practical training environments.",
        "focus": ["Windows/Linux server support and incident troubleshooting", "Active Directory, DNS/DHCP, Hyper-V, Docker", "Network diagnostics, monitoring, backups, runbooks"],
        "letter": "I am applying because the role’s focus on Windows/Linux support, system administration, AD, DNS/DHCP, monitoring, and incident resolution maps directly to my infrastructure work. At AnaxisTech, I support production Windows Server services and Docker-based GitLab, troubleshoot LAN/WAN and access issues, and maintain backup, recovery, and operational documentation. I am comfortable collaborating across infrastructure and development teams and can contribute from day one in a hands-on support environment.",
    },
    {
        "slug": "03_support_engineer_synoptek_ahmedabad",
        "title": "Support Engineer I",
        "company": "Synoptek",
        "location": "Ahmedabad, Gujarat",
        "url": "https://www.shine.com/jobs/support-engineer-i/synoptek/19167161",
        "summary": "Support-focused Network & Infrastructure Engineer experienced in Active Directory administration, DNS/DHCP troubleshooting, Windows support, LAN/WAN diagnostics, user access, server operations, and ticket-quality technical documentation.",
        "focus": ["Active Directory, passwords/permissions, Windows Server", "DNS/DHCP and network-connectivity troubleshooting", "End-user technical support, incident tracking, documentation"],
        "letter": "Synoptek’s Support Engineer I opportunity is a strong match for my background in Windows, Active Directory, DNS/DHCP, and network troubleshooting. I currently administer Windows Server 2025 services, maintain GPO security baselines, troubleshoot connectivity and SFTP/FTP issues, and provide infrastructure support to a developer team. I would bring a customer-focused support approach together with practical server and network depth, allowing me to resolve incidents thoroughly and document solutions clearly.",
    },
    {
        "slug": "04_network_security_engineer_aeromesh_ahmedabad",
        "title": "Senior Network & Security Engineer",
        "company": "Aeromesh Systems",
        "location": "Ahmedabad, Gujarat",
        "url": "https://www.shine.com/jobs/senior-network-security-engineer/aeromesh-systems/19184316",
        "summary": "Network & Infrastructure Engineer with practical routing/switching, VLAN segmentation, RRAS multi-ISP redundancy, firewall hardening, VPN/network security concepts, Windows Server, Docker, backup/DR, and infrastructure documentation experience.",
        "focus": ["VLANs, routing/switching, ACLs, NAT/PAT, RRAS", "Firewall hardening, network segmentation, VPN concepts", "Datacenter/server integration, backups, HA-minded documentation"],
        "letter": "I am interested in Aeromesh Systems because your work in enterprise switching, VLANs, routing, network segmentation, firewall systems, resilience, and datacenter connectivity is aligned with the infrastructure work I have already performed. I configured multi-ISP RRAS redundancy, Windows Server/AD services, firewall and port-forwarding rules, and automated backup/DR processes in production. My CCNA-focused training and Wazuh lab work reinforce this foundation. While my current title is Network & Infrastructure Engineer rather than a senior title, I am ready to contribute hands-on and grow quickly in an enterprise network-security environment.",
    },
    {
        "slug": "05_cisco_network_engineer_sndk_ahmedabad",
        "title": "Cisco Network Engineer",
        "company": "SNDK Corp",
        "location": "Ahmedabad, Gujarat",
        "url": "https://www.shine.com/jobs/cisco-network-engineer/sndk-corp/19206495",
        "summary": "Network Engineer with CCNA-focused practical training and production infrastructure experience covering TCP/IP, VLANs, routing/switching, OSPF, ACLs, NAT/PAT, RRAS, DNS/DHCP, troubleshooting, documentation, and security hardening.",
        "focus": ["TCP/IP, subnetting, VLANs, OSPF, ACLs, NAT/PAT, STP", "Routing/switching troubleshooting and network diagnostics", "Network security, RRAS redundancy, topology/runbook documentation"],
        "letter": "The Cisco Network Engineer role is closely aligned with my practical networking foundation and current production infrastructure work. My full-time CCNA routing-and-switching training covered subnetting, VLANs, OSPF, ACLs, NAT/PAT, and STP through daily labs; I now apply these skills while supporting a multi-ISP RRAS environment, Windows Server networking, and production troubleshooting. I am particularly interested in developing deeper Cisco implementation expertise while contributing disciplined troubleshooting, security awareness, and clear operational documentation.",
    },
    {
        "slug": "06_network_engineer_l2_gift_city",
        "title": "Network Engineer L2",
        "company": "Talent Vision Services",
        "location": "GIFT City / Ahmedabad, Gujarat",
        "url": "https://www.shine.com/jobs/network-engineer-l2-gift-city-ahmedabad/talent-vision-services/19098154",
        "summary": "Network & Infrastructure Engineer with production experience in network troubleshooting, multi-ISP RRAS redundancy, Windows Server, firewall hardening, Docker-hosted services, automated backup/DR, and infrastructure documentation; grounded in CCNA routing and switching practice.",
        "focus": ["Network infrastructure, routers/switches, troubleshooting", "Firewall hardening, security, network performance, RRAS", "Network upgrades, backups, documentation, incident support"],
        "letter": "I am applying for the Network Engineer L2 role because its focus on network infrastructure, routers, switches, firewalls, troubleshooting, security, performance, and documentation directly reflects my current work. At AnaxisTech, I implemented multi-ISP RRAS redundancy, supported Windows Server networking and Docker-hosted GitLab, performed firewall/VPS hardening, and maintained topology diagrams and runbooks. I am based in Ahmedabad and willing to work in GIFT City. I would be excited to bring a practical infrastructure mindset and strong learning velocity to the role.",
    },
]

def setup_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.55)
    sec.left_margin = sec.right_margin = Inches(0.65)
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(9.5)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def build_resume(job):
    doc = Document(); setup_doc(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SUBHAM HIRANI'); r.bold = True; r.font.size = Pt(17)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(job['title'] + ' | Network & Infrastructure Engineer').bold = True
    p = doc.add_paragraph(CONTACT); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph(job['summary'])
    doc.add_heading('ROLE-ALIGNED SKILLS', level=1)
    add_bullets(doc, job['focus'])
    doc.add_paragraph(BASE_SKILLS)
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    for heading, bullets in EXPERIENCE:
        p = doc.add_paragraph(); p.add_run(heading).bold = True
        add_bullets(doc, bullets)
    doc.add_heading('SELECTED PROJECTS', level=1)
    add_bullets(doc, [
        'Self-Hosted GitLab CE on Docker: deployed persistent-volume GitLab, reverse proxy, authentication, and tested backup/restore workflows.',
        'Centralized Automated Backup & DR System: designed scheduled backups with execution logging and integrity verification.',
        'On-Premises Enterprise Infrastructure: designed a 3-VM, four-VLAN architecture with AD, GPO baseline, Gitea, Restic, and Uptime Kuma.',
        'Wazuh SIEM Security Lab: analyzed alerts from simulated brute-force, port-scan, and privilege-escalation activity; documented triage steps.',
    ])
    doc.add_heading('EDUCATION & TRAINING', level=1)
    doc.add_paragraph('B.Tech in Computer Engineering, Sal College of Engineering (2022–2026)\nNetworking with Security — Tops Technologies: full-time practical training in CCNA, Windows Server, Linux, AWS, and cybersecurity.')
    return doc


def build_letter(job):
    doc = Document(); setup_doc(doc)
    doc.add_paragraph('Subham Hirani\n' + CONTACT)
    doc.add_paragraph('13 July 2026')
    doc.add_paragraph('Hiring Team\n' + job['company'])
    doc.add_paragraph('Subject: Application for ' + job['title'])
    doc.add_paragraph('Dear Hiring Team,')
    doc.add_paragraph('I am writing to apply for the ' + job['title'] + ' position in ' + job['location'] + '.')
    doc.add_paragraph(job['letter'])
    doc.add_paragraph('My resume is attached for your consideration. I would welcome the opportunity to discuss how my production infrastructure work and practical networking background can support your team.')
    doc.add_paragraph('Sincerely,\nSubham Hirani')
    return doc

for job in JOBS:
    d = OUT / job['slug']; d.mkdir(parents=True, exist_ok=True)
    build_resume(job).save(d / 'Subham_Hirani_Tailored_CV.docx')
    build_letter(job).save(d / 'Subham_Hirani_Cover_Letter.docx')
    (d / 'application_details.txt').write_text(
        f"Role: {job['title']}\nCompany: {job['company']}\nLocation: {job['location']}\nApplication URL: {job['url']}\n\n"
        "Submission fields still required from candidate: current CTC, expected CTC, notice period, work authorization, relocation confirmation, and portal login/OTP handling.\n",
        encoding='utf-8',
    )
print(f'Created {len(JOBS)} tailored CVs, {len(JOBS)} cover letters, and six application-detail files under {OUT}')
